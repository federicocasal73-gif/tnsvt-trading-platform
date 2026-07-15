#!/usr/bin/env python3
"""Convert all .md architecture docs to .docx and .pdf"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

BASE = Path(__file__).parent
DOCS_DIR = BASE / "docs"
WORD_DIR = BASE / "word"
PDF_DIR = BASE / "pdf"

WORD_DIR.mkdir(exist_ok=True)
PDF_DIR.mkdir(exist_ok=True)


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert Markdown to Word document with formatting."""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    
    for line in lines:
        stripped = line.rstrip('\n')
        
        if stripped.startswith('```'):
            if in_code_block:
                if code_lines:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(40, 40, 40)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(stripped)
            continue
        
        if stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            h = doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith('| ') and '|' in stripped[2:]:
            doc.add_paragraph(stripped, style='List Bullet')
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
        elif stripped.strip() == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    
    doc.save(str(docx_path))
    return True


class ArchitecturePDF(FPDF):
    """Custom PDF class for architecture docs."""
    
    def header(self):
        self.set_font('Helvetica', 'B', 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, 'TNSVT V2 - Architecture Documentation', align='L')
        self.ln(10)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')
    
    def chapter_title(self, title):
        self.set_font('Helvetica', 'B', 16)
        self.set_text_color(30, 60, 120)
        self.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(30, 60, 120)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(6)
    
    def section_title(self, title):
        self.set_font('Helvetica', 'B', 13)
        self.set_text_color(40, 80, 140)
        self.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
    
    def subsection_title(self, title):
        self.set_font('Helvetica', 'B', 11)
        self.set_text_color(60, 60, 60)
        self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)
    
    def body_text(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(30, 30, 30)
        self.multi_cell(0, 6, text)
        self.ln(2)
    
    def code_block(self, code):
        self.set_font('Courier', '', 8)
        self.set_fill_color(240, 240, 240)
        self.set_text_color(40, 40, 40)
        for line in code.split('\n'):
            self.cell(0, 4.5, '  ' + line, fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(3)
    
    def bullet_item(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(30, 30, 30)
        self.cell(8, 6, chr(8226))
        self.multi_cell(0, 6, text)
        self.ln(1)


def md_to_pdf(md_path: Path, pdf_path: Path):
    """Convert Markdown to PDF with formatting."""
    pdf = ArchitecturePDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    
    for line in lines:
        stripped = line.rstrip('\n')
        
        if stripped.startswith('```'):
            if in_code_block:
                if code_lines:
                    pdf.code_block('\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(stripped)
            continue
        
        if stripped.startswith('# ') and not stripped.startswith('## '):
            pdf.chapter_title(stripped[2:])
        elif stripped.startswith('## '):
            pdf.section_title(stripped[3:])
        elif stripped.startswith('### '):
            pdf.subsection_title(stripped[4:])
        elif stripped.startswith('#### '):
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 7, stripped[5:], new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[2:])
            pdf.bullet_item(text)
        elif stripped.startswith('> '):
            pdf.set_font('Helvetica', 'I', 9)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(0, 5, stripped[2:])
            pdf.ln(2)
        elif stripped.startswith('|'):
            pdf.set_font('Courier', '', 7)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 4.5, stripped[:190], new_x="LMARGIN", new_y="NEXT")
        elif stripped.strip() == '':
            pdf.ln(2)
        else:
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
            text = text[:190]
            pdf.body_text(text)
    
    pdf.output(str(pdf_path))
    return True


def main():
    md_files = sorted(DOCS_DIR.glob("*.md"))
    print(f"Found {len(md_files)} markdown files\n")
    
    for md_file in md_files:
        name = md_file.stem
        
        docx_path = WORD_DIR / f"{name}.docx"
        print(f"Converting {md_file.name} -> DOCX...", end=" ")
        try:
            md_to_docx(md_file, docx_path)
            print(f"OK ({docx_path.stat().st_size // 1024} KB)")
        except Exception as e:
            print(f"ERROR: {e}")
        
        pdf_path = PDF_DIR / f"{name}.pdf"
        print(f"Converting {md_file.name} -> PDF...", end=" ")
        try:
            md_to_pdf(md_file, pdf_path)
            print(f"OK ({pdf_path.stat().st_size // 1024} KB)")
        except Exception as e:
            print(f"ERROR: {e}")
    
    print(f"\nDone! Files in:\n  {WORD_DIR}\n  {PDF_DIR}")


if __name__ == "__main__":
    main()
